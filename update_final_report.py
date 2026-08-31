from pathlib import Path
from shutil import copy2

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


SOURCE = Path(r"C:\Users\SCSM11\Downloads\FinalReport.docx")
OUTPUT = Path(r"C:\Users\SCSM11\source\repos\SPC\FinalReport_Updated.docx")


def replace_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def insert_after(paragraph, text, style="normal"):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    new_paragraph = Paragraph(element, paragraph._parent)
    try:
        new_paragraph.style = style
    except KeyError:
        new_paragraph.style = "normal"
    new_paragraph.add_run(text)
    return new_paragraph


def add_after(paragraph, items):
    current = paragraph
    for text, style in items:
        current = insert_after(current, text, style)
    return current


def paragraph_by_text(doc, text, occurrence=0):
    matches = [p for p in doc.paragraphs if p.text.strip() == text]
    if len(matches) <= occurrence:
        raise ValueError(f"Paragraph not found: {text}")
    return matches[occurrence]


def delete_from(paragraph):
    element = paragraph._element
    parent = element.getparent()
    while element is not None:
        next_element = element.getnext()
        # The final section-properties element carries the document's page
        # setup and must remain in the Word body.
        if element.tag.endswith("}sectPr"):
            break
        parent.remove(element)
        element = next_element


copy2(SOURCE, OUTPUT)
doc = Document(OUTPUT)

# Keep the high-level overview consistent with the added Billing/Payment area.
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("The system is designed around several major functional areas"):
        replace_text(paragraph, "The system is designed around several major functional areas, including Customer Management, Room Management, Price Management, Room Booking, Billing/Payment, and Reporting. The Customer Management function allows customer records to be added, edited, deleted, displayed, and searched. The Room Management function allows users to view room details, check available rooms, update room status, and search for specific rooms. The Price Management function allows room prices and membership discounts to be viewed and updated. The Room Booking function manages room selection, customer verification, booking dates, check-in, check-out, cancellation, searching, and booking records. The Billing/Payment function calculates booking charges, records successful payments, and generates receipts. The Reporting function generates summary information such as total bookings, total customers, completed payments, and total revenue.")
    elif paragraph.text.startswith("The scope of the Hotel Reservation System covers"):
        replace_text(paragraph, "The scope of the Hotel Reservation System covers the main operational functions required for managing hotel reservations. The system provides six main functional areas: Customer Management, Room Management, Price Management, Room Booking, Billing/Payment, and Reporting.")
    elif paragraph.text.startswith("The program begins execution from the main() function"):
        replace_text(paragraph, "The program begins execution from the main() function. Before displaying the main menu, the system initializes the hotel room records through the initializeRooms() function. The user can then access the main functional areas of the system: Customer Management, Room Management, Price Management, Room Booking, Billing/Payment, and Reporting.")
    elif paragraph.text.strip() == "The five main modules are:":
        replace_text(paragraph, "The six main modules are:")
    elif paragraph.text.strip() == "Reporting" and paragraph._p.getprevious() is not None:
        # Only the Section 3 module list needs the extra Billing/Payment item.
        previous = Paragraph(paragraph._p.getprevious(), paragraph._parent)
        if previous.text.strip() == "Room Booking":
            insert_after(previous, "Billing/Payment", "normal")

# Section 2: add an explicit Billing/Payment functional area.
replace_text(paragraph_by_text(doc, "The system consists of five main functional areas:"), "The system consists of six main functional areas:")

overview_link = paragraph_by_text(doc, "The functional areas are integrated through shared data. Customer information is used during room booking, room information is used to determine the selected room and price and booking information is used during payment processing and reporting. This integration allows the different parts of the system to work together as a complete application. This assignment requires individual modules to interact correctly and share data within the overall system.")
replace_text(overview_link, "The functional areas are integrated through shared data. Customer information is used during room booking and payment calculation, room information is used to determine the selected room and price, booking information is used during billing and payment processing, and booking and payment information are used to generate reports. This integration allows the different parts of the system to work together as a complete application. This assignment requires individual modules to interact correctly and share data within the overall system.")

reporting_heading = paragraph_by_text(doc, "2.4.5 Reporting")
replace_text(reporting_heading, "2.4.6 Reporting")
billing_heading = insert_after(paragraph_by_text(doc, "The module also provides check-in and check-out operations, booking cancellation, booking search, and display of all booking records. During check-out, the room is marked as available again."), "2.4.5 Billing/Payment", "Heading 3")
insert_after(billing_heading, "The Billing/Payment module calculates the total booking charge using the room type, number of nights, and applicable customer membership discount. It supports Cash, Credit Card, Debit Card, Online Banking, and E-Wallet payments. For cash payment, the system validates the amount received and calculates the change. After a confirmed payment, the system creates a payment record, marks the booking as paid, updates the customer's total spending, and displays a payment receipt.", "normal")

# Section 5: complete testing and output descriptions, keeping screenshot slots for the user.
intro = paragraph_by_text(doc, "5.1 Introduction Screen")
add_after(intro, [
    ("When the program starts, the system displays an introduction screen identifying the course, assignment, and semester. This screen provides a clear starting point before the user enters the Hotel Reservation System.", "normal"),
    ("Figure 5.1: Introduction screen.", "normal"),
    ("[Insert screenshot of the introduction screen here]", "normal"),
])

main_menu = paragraph_by_text(doc, "5.2 Main Menu")
add_after(main_menu, [
    ("After the introduction screen, the main menu provides access to Customer Management, Room Management, Price Management, Room Booking, Report, and Exit. The user enters a menu number to navigate to the selected function. Invalid menu input is rejected and the system prompts the user again.", "normal"),
    ("Figure 5.2: Main menu.", "normal"),
    ("[Insert screenshot of the main menu here]", "normal"),
])

customer = paragraph_by_text(doc, "5.3 Customer Management Entry")
add_after(customer, [
    ("The Customer Management menu allows hotel staff to add, display, edit, delete, and search customer records. Customer ID, name, IC number, phone number, membership type, and total spending are maintained for later use by the booking and billing functions.", "normal"),
    ("Figure 5.3: Customer Management menu.", "normal"),
    ("[Insert screenshot of the Customer Management menu here]", "normal"),
])

room = paragraph_by_text(doc, "5.4 Room Management Entry")
add_after(room, [
    ("The Room Management menu displays room information and availability. Staff can view all rooms, view available rooms, search for a room, and update a room status. Room data are used by the booking function to allocate an available room.", "normal"),
    ("Figure 5.4: Room Management menu.", "normal"),
    ("[Insert screenshot of the Room Management menu here]", "normal"),
])

booking = paragraph_by_text(doc, "5.5 Room Booking Entry")
add_after(booking, [
    ("The Room Booking menu supports creating reservations, checking in guests, checking out guests, cancelling bookings, searching booking records, and displaying all bookings. During a new booking, the system verifies customer and room information, validates the stay dates, calculates the number of nights, and transfers the confirmed booking to the payment process.", "normal"),
    ("Figure 5.5: Room Booking menu.", "normal"),
    ("[Insert screenshot of the Room Booking menu here]", "normal"),
])

price = paragraph_by_text(doc, "5.6 Price Management Entry")
add_after(price, [
    ("The Price Management menu allows staff to display and update room prices and membership discount rates. Changes to these values are used in subsequent booking and payment calculations, ensuring that the current price and eligible discount are applied.", "normal"),
    ("Figure 5.6: Price Management menu.", "normal"),
    ("[Insert screenshot of the Price Management menu here]", "normal"),
])

report = paragraph_by_text(doc, "5.7 Report Entry")
add_after(report, [
    ("The Report menu provides three reporting options: Search by Month, Search by Year, and Overall Report. The generated report displays the total number of non-cancelled bookings, unique customers who have booking records, completed payments, and total revenue from successful payments.", "normal"),
    ("Figure 5.7: Report menu and generated report.", "normal"),
    ("[Insert screenshot of the Report menu and one generated report here]", "normal"),
])

exit_menu = paragraph_by_text(doc, "5.8 Exit")
add_after(exit_menu, [
    ("When the user selects Exit from the main menu, the system displays an exit message and terminates the program safely. This completes the program flow without leaving the user in a menu loop.", "normal"),
    ("Figure 5.8: Exit message.", "normal"),
    ("[Insert screenshot of the exit message here]", "normal"),
])

# Section 10: conclusion content.
summary = paragraph_by_text(doc, "10.1 Summary of the System")
add_after(summary, [
    ("The Hotel Reservation System is a console-based C++ application developed to support the main daily operations of a hotel. It integrates customer management, room and inventory management, price and discount management, booking and scheduling, billing and payment processing, and management reporting into one menu-driven system.", "normal"),
    ("The system stores customer, room, booking, payment, and report information in shared data structures. This allows information created in one module to be used by other modules. For example, customer membership and room price data are used to calculate payment charges, while booking and payment records are used to generate management reports.", "normal"),
])

achievement = paragraph_by_text(doc, "10.2 Achievement of Objectives")
add_after(achievement, [
    ("The main objectives of the system have been achieved:", "normal"),
    ("Manage customer records with validation for important customer details.", "List Bullet"),
    ("Maintain room information, room availability, room prices, and membership discounts.", "List Bullet"),
    ("Create, search, cancel, check in, and check out hotel bookings.", "List Bullet"),
    ("Calculate booking charges, apply membership discounts, process payments, and generate receipts.", "List Bullet"),
    ("Generate monthly, yearly, and overall reports for booking and revenue monitoring.", "List Bullet"),
    ("Apply modular programming, structures, arrays, vectors, functions, validation, selection, repetition, and formatted output.", "List Bullet"),
])

performance = paragraph_by_text(doc, "10.3 Overall System Performance")
add_after(performance, [
    ("Overall, the system performs the intended hotel-reservation workflow in a structured manner. The menu-driven design makes the functions easy to access, while validation helps to reduce common input errors. The shared data structures allow the modules to communicate consistently, for example when payment updates a booking record and reporting uses the payment data to calculate revenue.", "normal"),
    ("The system is suitable for demonstrating the required Systems and Programming Concepts in a small hotel environment. It provides clear console output, handles common invalid inputs without crashing, and gives staff an organised workflow from customer registration to reporting.", "normal"),
])

future = paragraph_by_text(doc, "10.4 Future Improvement")
add_after(future, [
    ("The following improvements can be considered for future development:", "normal"),
    ("Add file or database storage so that customer, booking, payment, and report records remain available after the program closes.", "List Bullet"),
    ("Add user login and role-based access control to protect staff functions and sensitive customer information.", "List Bullet"),
    ("Prevent a booking that has already been paid from being processed for payment again.", "List Bullet"),
    ("Add more report filters, such as custom date ranges, room-type revenue, occupancy rate, and exportable report files.", "List Bullet"),
    ("Develop a graphical user interface to improve usability for hotel staff.", "List Bullet"),
    ("Integrate a secure real-world payment gateway for card, online banking, and e-wallet transactions.", "List Bullet"),
])

# Remove an obsolete draft appended after the conclusion; it is not part of the final report.
obsolete = paragraph_by_text(doc, "Module Technical Description")
delete_from(obsolete)

doc.save(OUTPUT)
print(OUTPUT)
